"""Reporting tools: Markdown / PDF / DOCX export, chart, timeline, notification.

Heavy rendering libraries (reportlab, python-docx, matplotlib) are imported lazily inside
each tool so importing this module never requires them. Binary tools return ``bytes``.
"""

from __future__ import annotations

from pydantic import BaseModel, Field

from app.domain.exceptions import ToolExecutionError
from app.tools.base import BaseTool


# ── Structured report input ──────────────────────────────────
class ReportSections(BaseModel):
    title: str
    executive_summary: str = ""
    incident_timeline: list[str] = Field(default_factory=list)
    evidence: list[str] = Field(default_factory=list)
    root_cause: str = ""
    business_impact: str = ""
    technical_impact: str = ""
    recommendations: list[str] = Field(default_factory=list)
    action_items: list[str] = Field(default_factory=list)
    lessons_learned: str = ""


def render_markdown(sections: ReportSections) -> str:
    def bullets(items: list[str]) -> str:
        return "\n".join(f"- {item}" for item in items) if items else "_None_"

    return "\n".join(
        [
            f"# {sections.title}",
            "",
            "## Executive Summary",
            sections.executive_summary or "_N/A_",
            "",
            "## Incident Timeline",
            bullets(sections.incident_timeline),
            "",
            "## Evidence",
            bullets(sections.evidence),
            "",
            "## Root Cause",
            sections.root_cause or "_Undetermined_",
            "",
            "## Business Impact",
            sections.business_impact or "_N/A_",
            "",
            "## Technical Impact",
            sections.technical_impact or "_N/A_",
            "",
            "## Recommendations",
            bullets(sections.recommendations),
            "",
            "## Action Items",
            bullets(sections.action_items),
            "",
            "## Lessons Learned",
            sections.lessons_learned or "_N/A_",
            "",
        ]
    )


class MarkdownReportTool(BaseTool):
    name = "markdown_report"
    description = "Render a structured incident report as Markdown."
    args_schema = ReportSections

    async def arun(self, **kwargs) -> str:
        return render_markdown(ReportSections(**kwargs))


class ExportInput(BaseModel):
    title: str
    content: str = Field(description="Report body (Markdown or plain text).")


class PdfReportTool(BaseTool):
    name = "pdf_report"
    description = "Render a report to PDF bytes."
    args_schema = ExportInput

    async def arun(self, *, title: str, content: str) -> bytes:
        import io

        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER, title=title)
        styles = getSampleStyleSheet()
        flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                flow.append(Spacer(1, 6))
            elif stripped.startswith("# "):
                flow.append(Paragraph(stripped[2:], styles["Heading1"]))
            elif stripped.startswith("## "):
                flow.append(Paragraph(stripped[3:], styles["Heading2"]))
            else:
                flow.append(Paragraph(stripped.replace("- ", "• "), styles["BodyText"]))
        doc.build(flow)
        return buffer.getvalue()


class WordReportTool(BaseTool):
    name = "docx_report"
    description = "Render a report to DOCX bytes."
    args_schema = ExportInput

    async def arun(self, *, title: str, content: str) -> bytes:
        import io

        from docx import Document

        document = Document()
        document.add_heading(title, level=0)
        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("# "):
                document.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                document.add_heading(stripped[3:], level=2)
            elif stripped.startswith("- "):
                document.add_paragraph(stripped[2:], style="List Bullet")
            else:
                document.add_paragraph(stripped)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()


class ChartInput(BaseModel):
    title: str
    x_label: str = "time"
    y_label: str = "value"
    x: list[str] = Field(description="X-axis category/time labels.")
    y: list[float] = Field(description="Y-axis values, same length as x.")


class ChartGenerationTool(BaseTool):
    name = "chart_generation"
    description = "Render a simple line chart of a metric series to PNG bytes."
    args_schema = ChartInput

    async def arun(self, *, title: str, x: list[str], y: list[float], x_label: str = "time", y_label: str = "value") -> bytes:
        if len(x) != len(y):
            raise ToolExecutionError("x and y must have the same length.")
        import io

        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(x, y, marker="o")
        ax.set_title(title)
        ax.set_xlabel(x_label)
        ax.set_ylabel(y_label)
        fig.autofmt_xdate(rotation=30)
        fig.tight_layout()
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", dpi=100)
        plt.close(fig)
        return buffer.getvalue()


class TimelineInput(BaseModel):
    events: list[dict] = Field(description="Events with at least 'timestamp' and 'label'.")


class TimelineGeneratorTool(BaseTool):
    name = "timeline_generator"
    description = "Sort and normalize incident events into a chronological timeline."
    args_schema = TimelineInput

    async def arun(self, *, events: list[dict]) -> list[dict]:
        normalized = [
            {
                "timestamp": e.get("timestamp", ""),
                "label": e.get("label", "event"),
                "detail": e.get("detail", ""),
                "source": e.get("source", "unknown"),
            }
            for e in events
        ]
        return sorted(normalized, key=lambda e: e["timestamp"])


class NotificationInput(BaseModel):
    incident_id: str
    status: str
    probable_cause: str
    next_steps: str
    channel: str = "slack"
    audience: str = "#incidents"


class NotificationGeneratorTool(BaseTool):
    name = "notification_generator"
    description = "Format a concise stakeholder notification message."
    args_schema = NotificationInput

    async def arun(
        self,
        *,
        incident_id: str,
        status: str,
        probable_cause: str,
        next_steps: str,
        channel: str = "slack",
        audience: str = "#incidents",
    ) -> dict:
        message = (
            f":rotating_light: *Incident {incident_id}* — status: *{status}*\n"
            f"*Probable cause:* {probable_cause}\n"
            f"*Next steps:* {next_steps}"
        )
        return {"channel": channel, "audience": audience, "message": message}
